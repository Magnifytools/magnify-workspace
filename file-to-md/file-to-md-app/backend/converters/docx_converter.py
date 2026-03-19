import io
from docx import Document


def convert_docx(content: bytes, filename: str = "") -> str:
    doc = Document(io.BytesIO(content))
    md_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_parts.append("")
            continue

        style = para.style.name.lower() if para.style else ""

        if "heading 1" in style:
            md_parts.append(f"# {text}")
        elif "heading 2" in style:
            md_parts.append(f"## {text}")
        elif "heading 3" in style:
            md_parts.append(f"### {text}")
        elif "heading 4" in style:
            md_parts.append(f"#### {text}")
        elif "list" in style:
            md_parts.append(f"- {text}")
        else:
            # Handle bold/italic runs
            formatted = _format_runs(para.runs)
            md_parts.append(formatted if formatted.strip() else text)

    # Convert tables
    for table in doc.tables:
        md_parts.append("")
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        md_parts.append("| " + " | ".join(headers) + " |")
        md_parts.append("| " + " | ".join("---" for _ in headers) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            md_parts.append("| " + " | ".join(cells) + " |")
        md_parts.append("")

    return "\n".join(md_parts).strip()


def _format_runs(runs):
    parts = []
    for run in runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)
